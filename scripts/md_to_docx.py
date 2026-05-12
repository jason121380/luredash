"""One-shot: convert PRIVACY.md to PRIVACY.docx with reasonable styling.

Keeps headings, paragraphs, lists, simple tables, bold/italic inlines.
Skips fancy markdown (blockquotes get rendered as italics; code spans
get monospace font; horizontal rules become a thin centred line).

Not a general-purpose md->docx converter — only handles the subset of
markdown that PRIVACY.md actually uses.
"""

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

SRC = Path("PRIVACY.md")
DST = Path("PRIVACY.docx")

doc = Document()

# Set default font to a CJK-friendly stack
style = doc.styles["Normal"]
style.font.name = "Microsoft JhengHei"
style.font.size = Pt(11)
rpr = style.element.get_or_add_rPr()
rfonts = OxmlElement("w:rFonts")
rfonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
rfonts.set(qn("w:ascii"), "Microsoft JhengHei")
rfonts.set(qn("w:hAnsi"), "Microsoft JhengHei")
rpr.append(rfonts)

INLINE_BOLD = re.compile(r"\*\*([^*]+)\*\*")
INLINE_CODE = re.compile(r"`([^`]+)`")


def add_inline(paragraph, text: str):
    """Parse a small subset of inline markdown into runs."""
    pos = 0
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(1)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


lines = SRC.read_text(encoding="utf-8").splitlines()
i = 0
while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    # Skip empty lines
    if not stripped:
        i += 1
        continue

    # Heading
    if stripped.startswith("#"):
        m = re.match(r"^(#+)\s*(.*)$", stripped)
        level = min(len(m.group(1)), 4)
        text = m.group(2)
        h = doc.add_heading(level=level)
        add_inline(h, text)
        i += 1
        continue

    # Horizontal rule
    if stripped == "---":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("─" * 40).font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
        i += 1
        continue

    # Blockquote (use italic style)
    if stripped.startswith(">"):
        block_lines = []
        while i < len(lines) and lines[i].strip().startswith(">"):
            block_lines.append(re.sub(r"^>\s?", "", lines[i].strip()))
            i += 1
        joined = " ".join(b for b in block_lines if b)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        run = p.add_run(joined)
        run.italic = True
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        continue

    # Table — leading `|` and the next line is `| --- | ...`
    if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s|:-]+\|$", lines[i + 1].strip()):
        header_cells = [c.strip() for c in stripped.strip("|").split("|")]
        i += 2  # skip header + divider
        rows = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
            rows.append(cells)
            i += 1
        table = doc.add_table(rows=1 + len(rows), cols=len(header_cells))
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for idx, cell_text in enumerate(header_cells):
            run = hdr[idx].paragraphs[0].add_run(cell_text)
            run.bold = True
        for r_idx, row in enumerate(rows):
            for c_idx, cell_text in enumerate(row[: len(header_cells)]):
                add_inline(table.rows[r_idx + 1].cells[c_idx].paragraphs[0], cell_text)
        continue

    # Unordered list
    if stripped.startswith(("- ", "* ")):
        item_text = stripped[2:]
        p = doc.add_paragraph(style="List Bullet")
        add_inline(p, item_text)
        i += 1
        continue

    # Ordered list
    if re.match(r"^\d+\.\s", stripped):
        item_text = re.sub(r"^\d+\.\s+", "", stripped)
        p = doc.add_paragraph(style="List Number")
        add_inline(p, item_text)
        i += 1
        continue

    # Plain paragraph
    p = doc.add_paragraph()
    add_inline(p, stripped)
    i += 1

doc.save(DST)
print(f"wrote {DST} ({DST.stat().st_size} bytes)")
